import sys
import re
import statistics
import string
import random
import os.path
import itertools
import copy
from collections import defaultdict, Counter
import numpy as np
import pandas as pd
import copy
import docx
import docx.shared
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
#import tensorflow as tf

yes_no_map = {
"0": "no",
"1": "yes"
}

# From https://stackoverflow.com/questions/47738013/how-to-rotate-text-in-table-cells
def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)

# Adapted from https://stackoverflow.com/questions/40596518/writing-a-python-pandas-dataframe-to-word-document
def df_to_docx(df, header_bold=True):
    document = docx.Document()
    table = document.add_table(rows=(df.shape[0]+1), cols=df.shape[1])
    table.style = 'Table Grid'
    for j in range(df.shape[-1]):
        table.cell(0, j).text = df.columns[j]
        table.cell(0, j).paragraphs[0].runs[0].font.bold = header_bold
        table.cell(0, j).paragraphs[0].runs[0].font.size = docx.shared.Pt(6)


    table.rows[0].height = docx.shared.Inches(1)
    for i, column in enumerate(df) :
        for row in range(df.shape[0]):
            table.cell(row+1, i).text = str(df[column][row])
            table.cell(row+1, i).paragraphs[0].runs[0].font.size = docx.shared.Pt(6)
            if i == 0:
                table.cell(row+1, i).width = docx.shared.Inches(4)
    return(table, document)

def get_VCF_format_map(format_field, data_field):
    m = {}
    format_list = format_field.split(":")
    data_list = data_field.split(":")
    for i in range(0, len(format_list)):
        m[format_list[i]] = data_list[i]
    return(m)

# Based on Egor's code from ExpansionHunter Denovo - from source file IrrFinder.cpp
def MinimialUnitUnderShift(unit):
    minimal_unit = unit
    double_unit = unit + unit
    for index in range(0, len(unit)):
        current_unit = double_unit[index:index+len(unit)]
        if current_unit < minimal_unit:
            minimal_unit = current_unit
    return minimal_unit

# Based on Egor's code from ExpansionHunter Denovo - from source file IrrFinder.cpp
def ComputeCanonicalRepeatUnit(unit):
    minimal_unit = MinimialUnitUnderShift(unit)
    unit_rc = reverse_complement(unit)
    minimal_unit_rc = MinimialUnitUnderShift(unit_rc)

    if minimal_unit_rc < minimal_unit:
        return minimal_unit_rc
    return minimal_unit

# Check if two intervals overlap
def overlaps(x1, x2, y1, y2):
    return(x1 <= y2 and y1 <= x2)

def get_repeat_header(repeat_filename, num_segments):
    repeat_types = array_from_file(repeat_filename + ".repeat_types")
    return(["[D] Segment {:d} {}".format(segment_num, repeat_type) for segment_num in range(1, num_segments + 1) for repeat_type in repeat_types]) # "Segment 1 Gap", "Segment 1 Homopolymer", etc.

def get_repeat_features_from_segments(segments, repeat_filename, CNV_chr):
    min_start = min([s[0] for s in segments])
    repeat_types = array_from_file(repeat_filename + ".repeat_types")
    index_file = open(repeat_filename + ".idx")
    for index_file_line in index_file:
        index_chr, index_start, _, offset = index_file_line.rstrip("\n").replace("chr", "").split("\t")
        if CNV_chr == index_chr and min_start >= int(index_start):
            repeat_file_offset = int(offset)
    index_file.close()

    repeat_file = open(repeat_filename)
    repeat_features = []
    #Assigning repeat overlap values to each segment

    for segment_start, segment_end in segments:
        repeat_type_proportions = {}

        for repeat_type in repeat_types:
            repeat_type_proportions[repeat_type] = 0

        repeat_file.seek(repeat_file_offset)
        for line in repeat_file:
            repeat_chr, repeat_start, repeat_end, repeat_type = line.rstrip("\n").replace("chr", "").split("\t")
            repeat_start = int(repeat_start)
            repeat_end = int(repeat_end)
            if CNV_chr != repeat_chr:
                break
            elif repeat_start > segment_end: #if start of repeat is greater than end of CNV
                break
            elif segment_start <= repeat_end: #if start of CNV is greater than end of repeat
                overlap = get_overlap([segment_start, segment_end], [repeat_start, repeat_end])
                repeat_type_proportions[repeat_type] += overlap / (segment_end - segment_start + 1)

        for repeat_type in repeat_types:
            repeat_features.append("{:.3f}".format(repeat_type_proportions[repeat_type]))
    repeat_file.close()
    return(repeat_features)

def get_repeat_features_spanning_breakpoints(repeat_filename, CNV_chr, CNV_start, CNV_end, num_bp_spanning):
    CNV_chr = CNV_chr.replace("chr", "")
    CNV_start = int(CNV_start)
    CNV_end = int(CNV_end)

    segments = [[CNV_start - num_bp_spanning/2, CNV_start + num_bp_spanning/2], [CNV_end - num_bp_spanning/2, CNV_end + num_bp_spanning/2]]
    return(get_repeat_features_from_segments(segments, repeat_filename, CNV_chr))

def get_repeat_features_extended_segments(repeat_filename, CNV_chr, CNV_start, CNV_end, extend_proportion, num_segments):
    CNV_chr = CNV_chr.replace("chr", "")
    CNV_start = int(CNV_start)
    CNV_end = int(CNV_end)
    CNV_size = CNV_end - CNV_start + 1

    chr_size = get_hg19_chromosome_sizes()

    extend_size = CNV_size * extend_proportion

    extended_CNV_start = max(1, CNV_start - extend_size)
    extended_CNV_end = min(CNV_end + extend_size, chr_size[CNV_chr])
    extended_CNV_size = extended_CNV_end - extended_CNV_start + 1

    # Create a list of length num_segments, with each element being a two-element list containing the start and end coordinates of that segment
    segments = [[int(extended_CNV_start + extended_CNV_size * i/num_segments), int(extended_CNV_start + extended_CNV_size * (i+1)/num_segments) - 1] for i in range(num_segments)]

    return(get_repeat_features_from_segments(segments, repeat_filename, CNV_chr))

def size_bin_bp_to_kb(size_bin):
    reg = re.search(r'(\d+),([\d.]+)', size_bin)
    if reg:
        lower = reg.group(1)
        upper = reg.group(2)

        if upper == "...":
            final_upper = "..."
        else:
            uf = int(upper) / 1000
            if uf.is_integer():
                final_upper = "{:d} kb".format(int(uf))
            else:

                final_upper = "{:.1f} kb".format(uf)


        lf = int(lower) / 1000
        if lf.is_integer():
            final_lower = "{:d} kb".format(int(lf))
        else:
            final_lower = "{:.1f} kb".format(lf)

        return("[{},{})".format(final_lower, final_upper))
    else: # "Total" or similar
        return(size_bin)



def get_header_line(filename):
    f = open(filename)
    header = f.readline().rstrip("\n")
    f.close()
    return(header)

# Read in the file containing the training data
def prepare_data_for_tensorflow(filename):
    infile = open(filename)
    original_header = infile.readline().rstrip("\n").split("ORIGINAL_DATA_HEADER:")[1]
    data_raw = pd.read_csv(infile, sep="\t") # Read in data as pandas data frame
    instance_names = data_raw.iloc[:, 0].ravel() # Get the instance names
    labels_raw = data_raw.iloc[:, 1].ravel() # Make a numpy array of labels
    data = np.array(data_raw.iloc[:, 2:]) # Make a numpy array containing the training data (all rows, and all columns but the first and second, which contain the instance names and labels)
    return(data, labels_raw, instance_names, original_header)

# Convert a numpy array to one-hot format, where the index corresponding to a given class is 1, and all other elements are 0.
def convert_to_one_hot(labels):
    one_hot_array = np.zeros((len(labels), 2)) # Create numpy array with rows = number of labels and columns = number of classes (=2), with all elements 0
    index_row = np.arange(len(labels)) * 2 # Create a numpy array containing 0, 2, 4, 6, ... (number of elements = number of labels)
    one_hot_array.flat[[index_row + labels]] = 1 # Create a numpy array, where each element contains two elements. If there is a 1 in the first position, then the classification is 0 (not a CNV); if there is a 1 in the second position, then the classification is 1 (CNV)
    return(one_hot_array)

def get_VCF_format_dict(format, data):
    format_split = format.split(":")
    data_split = data.split(":")

    d = defaultdict(str)
    for i in range(0, len(format_split)):
        d[format_split[i]] = data_split[i]
    return(d)

def gene_start(exon_data): # data for a particular cds, as read in by read_cds_definition
    min_first = 999999999
    for exon in exon_data:
        start_pos = int(float(exon["start"]))
        if start_pos < min_first:
            min_first = start_pos
    return(min_first)

def gene_end(exon_data):   # data for a particular cds, as read in by read_cds_definition
    max_end = 0
    for exon in exon_data:
        end_pos = int(float(exon["end"]))
        if end_pos > max_end:
            max_end = end_pos
    return(max_end)

def read_cds_definition(cds_definition_filename="/hpf/largeprojects/tcagstor/tools/annotations/CNVannotationPipeline/V17.1/20171204/fixedCDS.tsv"):
    cds_definition_file = open(cds_definition_filename)
    cds_definition_file.readline() # Discard header line

    cds_definition = defaultdict(list)

    for line in cds_definition_file:
        chrom, start, end, cdsEgID, cdsSymbol  = line.rstrip("\n").split("\t")
        cds_definition[cdsSymbol].append( {"chr": chrom, "start": start, "end": end} )

    return(cds_definition)


def find_cumulative_val(v, total, percent):
    cumulative_sum = 0

    for i in sorted(list(v.keys())):
	    cumulative_sum = cumulative_sum + v[i]

	    if cumulative_sum > total * (percent / 100):
		    quartile_num = i
		    break

    return(quartile_num)

# Generator used so it doesn't use a large amount of memory when reading a big file
def read_table_as_dict(filename, return_header=True, skip_lines=0, delimiter="\t"):
    infile = file_or_stdin(filename)

    for i in range(0, skip_lines):
        yield(infile.readline().rstrip("\n"))

    header = infile.readline().rstrip("\n")
    column_names = header.split(delimiter)
    if return_header:
        yield(header)
    for l in infile:
        this_rec = {}
        fields = l.rstrip("\n").split(delimiter)
        for i in range(0, len(fields)):
            this_rec[column_names[i]] = fields[i]
        this_rec["_FULL_LINE"] = l.rstrip("\n")
        yield(this_rec)


def eprint(*args, **kwargs):
    print(*args, file=sys.stderr, **kwargs)

# https://stackoverflow.com/questions/29348345/declaring-a-multi-dimensional-dictionary-in-python
def nested_dict(num_dimensions, type):
    if num_dimensions == 1:
        return defaultdict(type)
    else:
        return defaultdict(lambda: nested_dict(num_dimensions - 1, type))

def get_DELLY_SR_PE(SV):
    fields = SV["all_caller_info"].split("|")[2].split(":")

    num_PE = int(fields[9])
    num_SR = int(fields[11])

    return(num_SR, num_PE)

def get_LUMPY_SR_PE_qual(SV):
    fields = SV["all_caller_info"].split("|")[2].split(":")

    num_PE = int(fields[2])
    num_SR = int(fields[3])
    qual = float(SV["all_caller_info"].split(";")[0].split("=")[1])

    return(num_SR, num_PE, qual)

def get_Manta_SR_PE(SV):
    field_names = SV["all_caller_info"].split("|")[1].split(":")
    field_values = SV["all_caller_info"].split("|")[2].split(":")

    if "SR" in field_names:
        SR_index = field_names.index("SR")
        num_SR = int(field_values[SR_index].split(",")[1])
    else:
        num_SR = 0

    if "PR" in field_names:
        PE_index = field_names.index("PR")
        num_PE = int(field_values[PE_index].split(",")[1])
    else:
        num_PE = 0

    return(num_SR, num_PE)

def get_DELLY_SR_PE_for_display(group, for_features=False):
    num_SR = "N/A"
    num_PE = "N/A"

    if for_features:
        num_SR = -1
        num_PE = -1

    for SV in group:
        if SV["caller"] == "DELLY":
            num_SR, num_PE = get_DELLY_SR_PE(SV)
    return(num_SR, num_PE)

def get_LUMPY_SR_PE_qual_for_display(group, for_features=False):
    num_SR = "N/A"
    num_PE = "N/A"
    qual = "N/A"

    if for_features:
        num_SR = -1
        num_PE = -1
        qual = -1

    for SV in group:
        if SV["caller"] == "LUMPY":
            num_SR, num_PE, qual = get_LUMPY_SR_PE_qual(SV)
    return(num_SR, num_PE, qual)


def get_Manta_SR_PE_PASS_for_display(group, for_features=False):
    num_SR = "N/A"
    num_PE = "N/A"
    PASS = "N/A"

    if for_features:
        num_SR = -1
        num_PE = -1
        PASS = -1

    for SV in group:
        if SV["caller"] == "Manta":
            num_SR, num_PE = get_Manta_SR_PE(SV)
            PASS = SV["important_caller_info"]
            if for_features:
                if PASS == "PASS":
                    PASS = 1
                else:
                    PASS = 0

    return(num_SR, num_PE, PASS)



def super_voodoo_lt300(group, average_read_depth=40, LUMPY_only_ratio=0.1, other_ratio=0.125): # 0.1 = 4/40, 0.125 = 5/40

    num_DELLY_SR = 0
    num_LUMPY_SR = 0
    num_Manta_SR = 0

    five_count = 0

    for SV in group:
        if SV["caller"] == "DELLY":
            num_DELLY_SR, _ = get_DELLY_SR_PE(SV)
        if SV["caller"] == "LUMPY":
            num_LUMPY_SR, _, _ = get_LUMPY_SR_PE_qual(SV)
        if SV["caller"] == "Manta":
            num_Manta_SR, _ = get_Manta_SR_PE(SV)

    for num in [num_DELLY_SR, num_LUMPY_SR, num_Manta_SR]:
        if num/average_read_depth >= other_ratio:
            five_count += 1

    return(five_count >= 2 or (num_DELLY_SR/average_read_depth >= other_ratio and num_LUMPY_SR/average_read_depth >= LUMPY_only_ratio))

def super_voodoo_lt300_no_DELLY(group, average_read_depth=40, other_ratio=0.125): # 0.125 = 5/40

    num_LUMPY_SR = 0
    num_Manta_SR = 0

    for SV in group:
        if SV["caller"] == "LUMPY":
            num_LUMPY_SR, _, _ = get_LUMPY_SR_PE_qual(SV)
        if SV["caller"] == "Manta":
            num_Manta_SR, _ = get_Manta_SR_PE(SV)

    return(num_LUMPY_SR/average_read_depth >= other_ratio and num_Manta_SR/average_read_depth >= other_ratio)

def super_voodoo_gt300(group, average_read_depth=40, other_ratio=0.125):

    num_LUMPY_PE = 0
    num_Manta_PE = 0
    LUMPY_qual = 0

    Manta_pass = False

    for SV in group:
        if SV["caller"] == "LUMPY":
            _, num_LUMPY_PE, LUMPY_qual = get_LUMPY_SR_PE_qual(SV)
        if SV["caller"] == "Manta":
            _, num_Manta_PE = get_Manta_SR_PE(SV)
            if SV["important_caller_info"] == "PASS":
                Manta_pass = True

    return((Manta_pass and num_Manta_PE/average_read_depth >= other_ratio) or (LUMPY_qual >= 100 and num_LUMPY_PE/average_read_depth >= other_ratio))

def simple_super_voodoo_lt300(group, average_read_depth=40, LUMPY_only_ratio=0.1, other_ratio=0.125): # 0.1 = 4/40, 0.125 = 5/40

    num_DELLY_SR = 0
    num_LUMPY_SR = 0
    num_Manta_SR = 0

    for SV in group:
        if SV["caller"] == "DELLY":
            num_DELLY_SR, _ = get_DELLY_SR_PE(SV)
        if SV["caller"] == "LUMPY":
            num_LUMPY_SR, _, _ = get_LUMPY_SR_PE_qual(SV)
        if SV["caller"] == "Manta":
            num_Manta_SR, _ = get_Manta_SR_PE(SV)

    pass_count = 0
    if num_DELLY_SR/average_read_depth >= other_ratio:
        pass_count += 1
    if num_LUMPY_SR/average_read_depth >= LUMPY_only_ratio:
        pass_count += 1
    if num_Manta_SR/average_read_depth >= other_ratio:
        pass_count += 1

    return(pass_count >= 2)

def super_simple_super_voodoo_lt300(group, average_read_depth=40, LUMPY_only_ratio=0.1, other_ratio=0.125): # 0.1 = 4/40, 0.125 = 5/40

    num_DELLY_SR = 0
    num_LUMPY_SR = 0
    num_Manta_SR = 0

    for SV in group:
        if SV["caller"] == "DELLY":
            num_DELLY_SR, _ = get_DELLY_SR_PE(SV)
        if SV["caller"] == "LUMPY":
            num_LUMPY_SR, _, _ = get_LUMPY_SR_PE_qual(SV)
        if SV["caller"] == "Manta":
            num_Manta_SR, _ = get_Manta_SR_PE(SV)

    pass_count = 0
    if num_DELLY_SR/average_read_depth >= other_ratio:
        pass_count += 1
    if num_LUMPY_SR/average_read_depth >= other_ratio:
        pass_count += 1
    if num_Manta_SR/average_read_depth >= other_ratio:
        pass_count += 1

    return(pass_count >= 2)


#Definition of Super Simple Super Voodoo :
#<300 bp: 2 of Manta, LUMPY, and DELLY, where it passes Manta and DELLY and LUMPY with >= 5 split reads
#>300 bp: Manta PASS and >= 5 PR and/or LUMPY Qual >= 100 and >= 5PR
def super_simple_super_voodoo(group, average_read_depth=40, other_ratio=0.125):
    # Get median size
    size = median_size_of_SVs_in_group(group)
    if size < 300:
        return(super_simple_super_voodoo_lt300(group, average_read_depth=average_read_depth, other_ratio=other_ratio))
    else:
        return(super_voodoo_gt300(group, average_read_depth=average_read_depth, other_ratio=other_ratio))

#Definition of Super Voodoo Simple:
#<300 bp: 2 of Manta, LUMPY, and DELLY, where it passes Manta and DELLY with >= 5 split reads and LUMPY with >= 4 split reads.
#>300 bp: Manta PASS and >= 5 PR and/or LUMPY Qual >= 100 and >= 5PR
def simple_super_voodoo(group, average_read_depth=40, other_ratio=0.125):
    # Get median size
    size = median_size_of_SVs_in_group(group)

    if size < 300:
        return(simple_super_voodoo_lt300(group, average_read_depth=average_read_depth, other_ratio=other_ratio))
    else:
        return(super_voodoo_gt300(group, average_read_depth=average_read_depth, other_ratio=other_ratio))

#Definition of Super Voodoo:
#<300 bp: >= 5 split reads for any two of Manta, LUMPY and DELLY, unless it's LUMPY and DELLY, then DELLY >= 5 and LUMPY >= 4.
#>300 bp: Manta PASS and >= 5 PR and/or LUMPY Qual >= 100 and >= 5PR
def super_voodoo(group, average_read_depth=40, other_ratio=0.125):
    # Get median size
    size = median_size_of_SVs_in_group(group)

    if size < 300:
        return(super_voodoo_lt300(group, average_read_depth=average_read_depth, other_ratio=other_ratio))
    else:
        return(super_voodoo_gt300(group, average_read_depth=average_read_depth, other_ratio=other_ratio))

def super_voodoo_no_DELLY(group, average_read_depth=40, other_ratio=0.125):
    # Get median size
    size = median_size_of_SVs_in_group(group)

    if size < 300:
        return(super_voodoo_lt300_no_DELLY(group, average_read_depth=average_read_depth, other_ratio=other_ratio))
    else:
        return(super_voodoo_gt300(group, average_read_depth=average_read_depth, other_ratio=other_ratio))

def called_by_feature(group, caller):
    for SV in group:
        if SV["caller"] == caller:
            return(1)
    return(0)

def called_by(group, caller):
    for SV in group:
        if SV["caller"] == caller:
            return(True)
    return(False)

def MetaSV(group, average_read_depth=None, other_ratio=None):
    return(called_by(group, "MetaSV"))

def LUMPY_or_Manta(group, average_read_depth=None, other_ratio=None):
    return(called_by(group, "LUMPY") or called_by(group, "Manta"))

def LUMPY_or_Manta_with_ERDS(group, average_read_depth=None, other_ratio=None):
    return((called_by(group, "LUMPY") or called_by(group, "Manta")) and called_by(group, "ERDS"))

def super_voodoo_with_ERDS(group, average_read_depth=40, other_ratio=0.125):
    return(super_voodoo(group, average_read_depth=average_read_depth, other_ratio=other_ratio) and called_by(group, "ERDS"))

def one_method(group, average_read_depth=None, other_ratio=None):
    return(count_in_group(group, "prediction") >= 1)

def two_methods(group, average_read_depth=None, other_ratio=None):
    return(count_in_group(group, "prediction") >= 2)

def three_methods(group, average_read_depth=None, other_ratio=None):
    return(count_in_group(group, "prediction") >= 3)

def four_methods(group, average_read_depth=None, other_ratio=None):
    return(count_in_group(group, "prediction") >= 4)

def five_methods(group, average_read_depth=None, other_ratio=None):
    return(count_in_group(group, "prediction") >= 5)

def six_methods(group, average_read_depth=None, other_ratio=None):
    return(count_in_group(group, "prediction") >= 6)

# Two different ways of doing this: either make all possible pairwise comparisons and then merge them in decreasing order of pairwise identity (sort_by_overlap_first=True)
# or just merge them as you find them (sort_by_overlap_first=False). They are largely similar, although sort_by_overlap_first=True can actually give counterintuitive results.
def get_groups(CNVs_by_chromosome, methods_list, item_name_key, overlap_function, sort_by_overlap_first=False, ignore_type=False):
    chromosomes = get_chromosomes()
    groups = []
    merged_CNV_pointers = {}

    if sort_by_overlap_first:
        overlap_list = []
        for chrom in chromosomes:
            for i in range(0, len(methods_list) - 1):
                for k in range(0, len(CNVs_by_chromosome[methods_list[i]][chrom])):
                    for j in range(i+1, len(methods_list)):
                        for l in range(0, len(CNVs_by_chromosome[methods_list[j]][chrom])):
                            mything = {}

                            mything["first"] = CNVs_by_chromosome[methods_list[i]][chrom][k]
                            mything["second"] = CNVs_by_chromosome[methods_list[j]][chrom][l]

                            if overlap_function(mything["first"], mything["second"], ignore_type):
                                mything["overlap"] = avg_reciprocal_overlap(mything["first"], mything["second"])
                                overlap_list.append(mything)
                                CNVs_by_chromosome[methods_list[i]][chrom][k]["FOUND"] = True
                                CNVs_by_chromosome[methods_list[j]][chrom][l]["FOUND"] = True


        overlap_list = sorted(overlap_list, key=lambda k: k["overlap"], reverse=True)

        for overlap_item in overlap_list:
            #print("first is {}:{}-{} and second is {}:{}-{} and overlap is {}".format(overlap_item["first"]["chr"], overlap_item["first"]["start"], overlap_item["first"]["end"], overlap_item["second"]["chr"], overlap_item["second"]["start"], overlap_item["second"]["end"], overlap_item["overlap"]))

            if overlap_item["first"][item_name_key] in merged_CNV_pointers and overlap_item["second"][item_name_key] in merged_CNV_pointers:
                if merged_CNV_pointers[overlap_item["first"][item_name_key]] != merged_CNV_pointers[overlap_item["second"][item_name_key]]:
                    group1_names = [item[item_name_key] for item in merged_CNV_pointers[overlap_item["first"][item_name_key]]]
                    group2_names = [item[item_name_key] for item in merged_CNV_pointers[overlap_item["second"][item_name_key]]]
                continue
            elif overlap_item["first"][item_name_key] in merged_CNV_pointers: # i is already in a group, so just add j
                group = merged_CNV_pointers[overlap_item["first"][item_name_key]]
                group.append(overlap_item["second"])
                merged_CNV_pointers[overlap_item["second"][item_name_key]] = group
            elif overlap_item["second"][item_name_key] in merged_CNV_pointers: # j is already in a group, so just add i
                group = merged_CNV_pointers[overlap_item["second"][item_name_key]]
                group.append(overlap_item["first"])
                merged_CNV_pointers[overlap_item["first"][item_name_key]] = group
            else: # Neither one is in a group yet, so make a new group
                new_group = [overlap_item["first"], overlap_item["second"]]
                groups.append(new_group)
                merged_CNV_pointers[overlap_item["first"][item_name_key]] = new_group
                merged_CNV_pointers[overlap_item["second"][item_name_key]] = new_group

        for chrom in chromosomes:
            for i in range(0, len(methods_list)):
                for k in range(0, len(CNVs_by_chromosome[methods_list[i]][chrom])):
                    if "FOUND" not in CNVs_by_chromosome[methods_list[i]][chrom][k]:
                        new_group = [CNVs_by_chromosome[methods_list[i]][chrom][k]]
                        groups.append(new_group)
        return(groups)
    else:
        for chrom in chromosomes:
            for i in range(0, len(methods_list) - 1):
                for k in range(0, len(CNVs_by_chromosome[methods_list[i]][chrom])):
                    for j in range(i+1, len(methods_list)):
                        for l in range(0, len(CNVs_by_chromosome[methods_list[j]][chrom])):
                            if overlap_function(CNVs_by_chromosome[methods_list[i]][chrom][k], CNVs_by_chromosome[methods_list[j]][chrom][l], ignore_type):
                                if CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key] in merged_CNV_pointers and CNVs_by_chromosome[methods_list[j]][chrom][l][item_name_key] in merged_CNV_pointers:
                                    if merged_CNV_pointers[CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key]] != merged_CNV_pointers[CNVs_by_chromosome[methods_list[j]][chrom][l][item_name_key]]:
                                        group1_names = [item[item_name_key] for item in merged_CNV_pointers[CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key]]]
                                        group2_names = [item[item_name_key] for item in merged_CNV_pointers[CNVs_by_chromosome[methods_list[j]][chrom][l][item_name_key]]]
                                    continue
                                elif CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key] in merged_CNV_pointers: # i is already in a group, so just add j
                                    group = merged_CNV_pointers[CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key]]
                                    group.append(CNVs_by_chromosome[methods_list[j]][chrom][l])
                                    merged_CNV_pointers[CNVs_by_chromosome[methods_list[j]][chrom][l][item_name_key]] = group
                                elif CNVs_by_chromosome[methods_list[j]][chrom][l][item_name_key] in merged_CNV_pointers: # j is already in a group, so just add i
                                    group = merged_CNV_pointers[CNVs_by_chromosome[methods_list[j]][chrom][l][item_name_key]]
                                    group.append(CNVs_by_chromosome[methods_list[i]][chrom][k])
                                    merged_CNV_pointers[CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key]] = group
                                else: # Neither one is in a group yet, so make a new group
                                    new_group = [CNVs_by_chromosome[methods_list[i]][chrom][k], CNVs_by_chromosome[methods_list[j]][chrom][l]]
                                    groups.append(new_group)
                                    merged_CNV_pointers[CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key]] = new_group
                                    merged_CNV_pointers[CNVs_by_chromosome[methods_list[j]][chrom][l][item_name_key]] = new_group
                    if not CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key] in merged_CNV_pointers:
                        new_group = [CNVs_by_chromosome[methods_list[i]][chrom][k]]
                        groups.append(new_group)
                        merged_CNV_pointers[CNVs_by_chromosome[methods_list[i]][chrom][k][item_name_key]] = new_group


        ### Now we need to find unmatched CNVs_by_chromosome in the last tool ###
        last_i = len(methods_list) - 1
        for chrom in chromosomes:
            for k in range(0, len(CNVs_by_chromosome[methods_list[last_i]][chrom])):
                if not CNVs_by_chromosome[methods_list[last_i]][chrom][k][item_name_key] in merged_CNV_pointers:
                    new_group = [CNVs_by_chromosome[methods_list[last_i]][chrom][k]]
                    groups.append(new_group)
                    merged_CNV_pointers[CNVs_by_chromosome[methods_list[last_i]][chrom][k][item_name_key]] = new_group
        return(groups)


def get_groups_one_set(CNVs_by_chromosome, item_name_key, overlap_function, ignore_type=False):

    chromosomes = get_chromosomes()
    groups = []
    merged_CNV_pointers = {}

    for chrom in chromosomes:
        for k in range(0, len(CNVs_by_chromosome[chrom]) - 1):
            for l in range(k+1, len(CNVs_by_chromosome[chrom])):
                #print("CNV1: {}: CNV2:{} first".format(CNVs_by_chromosome[chrom][k][item_name_key], CNVs_by_chromosome[chrom][l][item_name_key]))
                if overlap_function(CNVs_by_chromosome[chrom][k], CNVs_by_chromosome[chrom][l], ignore_type):
                    #print("CNV1: {}: CNV2:{} overlapped".format(CNVs_by_chromosome[chrom][k][item_name_key], CNVs_by_chromosome[chrom][l][item_name_key]))
                    if CNVs_by_chromosome[chrom][k][item_name_key] in merged_CNV_pointers and CNVs_by_chromosome[chrom][l][item_name_key] in merged_CNV_pointers:
                        #print("CNV1: {}: CNV2:{} already joined".format(CNVs_by_chromosome[chrom][k][item_name_key], CNVs_by_chromosome[chrom][l][item_name_key]))
                        pass
                    elif CNVs_by_chromosome[chrom][k][item_name_key] in merged_CNV_pointers: # i is already in a group, so just add j
                        #print("CNV1: {}: CNV2:{} kkk".format(CNVs_by_chromosome[chrom][k][item_name_key], CNVs_by_chromosome[chrom][l][item_name_key]))
                        group = merged_CNV_pointers[CNVs_by_chromosome[chrom][k][item_name_key]]
                        group.append(CNVs_by_chromosome[chrom][l])
                        merged_CNV_pointers[CNVs_by_chromosome[chrom][l][item_name_key]] = group
                    elif CNVs_by_chromosome[chrom][l][item_name_key] in merged_CNV_pointers: # j is already in a group, so just add i
                        #print("CNV1: {}: CNV2:{} lll".format(CNVs_by_chromosome[chrom][k][item_name_key], CNVs_by_chromosome[chrom][l][item_name_key]))
                        group = merged_CNV_pointers[CNVs_by_chromosome[chrom][l][item_name_key]]
                        group.append(CNVs_by_chromosome[chrom][k])
                        merged_CNV_pointers[CNVs_by_chromosome[chrom][k][item_name_key]] = group
                    else: # Neither one is in a group yet, so make a new group
                        #print("CNV1: {}: CNV2:{} ggggg".format(CNVs_by_chromosome[chrom][k][item_name_key], CNVs_by_chromosome[chrom][l][item_name_key]))
                        new_group = [CNVs_by_chromosome[chrom][k], CNVs_by_chromosome[chrom][l]]
                        groups.append(new_group)
                        merged_CNV_pointers[CNVs_by_chromosome[chrom][k][item_name_key]] = new_group
                        merged_CNV_pointers[CNVs_by_chromosome[chrom][l][item_name_key]] = new_group
            if not CNVs_by_chromosome[chrom][k][item_name_key] in merged_CNV_pointers:
                new_group = [CNVs_by_chromosome[chrom][k]]
                groups.append(new_group)
                merged_CNV_pointers[CNVs_by_chromosome[chrom][k][item_name_key]] = new_group
    return(groups)


def get_submitted_ID(subject_id_ext, MSSNG_info):
    if subject_id_ext[-1] == "A" or subject_id_ext[-1] == "B" or subject_id_ext[-1] == "C" or subject_id_ext[-1] == "D":
        subject_id_ext = subject_id_ext[0:-1]

    subject_data = MSSNG_info.loc[MSSNG_info["SUBJECT_ID_EXT"] == subject_id_ext]

    subject_row_list = []

    for _, subject_row in subject_data.iterrows():
        subject_row_list.append(subject_row)

    if len(subject_row_list) == 1:
        return(subject_row_list[0]["SUBMITTED_ID"])
    else:
        best = 0
        best_subject_row = "ERROR: did not find SUBMITTED_ID"
        letter_map = {"A": 1, "B": 2, "C": 3, "D": 4}
        for subject_row in subject_row_list:
            if subject_row["SUBMITTED_ID"][-1] == "A" or subject_row["SUBMITTED_ID"][-1] == "B" or subject_row["SUBMITTED_ID"][-1] == "C" or subject_row["SUBMITTED_ID"][-1] == "D":
                if letter_map[subject_row["SUBMITTED_ID"][-1]] > best:
                    best = letter_map[subject_row["SUBMITTED_ID"][-1]]
                    best_subject_row = subject_row["SUBMITTED_ID"]

        return(best_subject_row)


def get_parent_to_children_dict(MSSNG_info):

    children_submitted_id_dict = defaultdict(list)

    for _, row in MSSNG_info.iterrows():
        if row["FATHER_ID"] != "0":
            father_submitted_id = get_submitted_ID(row["FATHER_ID"], MSSNG_info)
            children_submitted_id_dict[father_submitted_id].append(row["SUBMITTED_ID"])
        if row["MOTHER_ID"] != "0":
            mother_submitted_id = get_submitted_ID(row["MOTHER_ID"], MSSNG_info)
            children_submitted_id_dict[mother_submitted_id].append(row["SUBMITTED_ID"])

    return(children_submitted_id_dict)



def get_mother_submitted_ID(proband_submitted_ID, MSSNG_info):
    return(get_submitted_ID(MSSNG_info.loc[proband_submitted_ID]["MOTHER_ID"], MSSNG_info))

def get_father_submitted_ID(proband_submitted_ID, MSSNG_info):
    return(get_submitted_ID(MSSNG_info.loc[proband_submitted_ID]["FATHER_ID"], MSSNG_info))


def get_variant_str(VCF_rec):
    return(":".join([VCF_rec[0], VCF_rec[1], VCF_rec[3], VCF_rec[4]])) # chr:pos:ref:alt

# Given a list of VCF variants returned by tabix Python module, make a dictionary for which a particular
# key exists if the sample has a particular variant. Makes it easy to look up whether a particular variant is found in this individual
def get_variant_dict_from_VCF_recs(VCF_recs):
    variant_dict = {}
    for VCF_rec in VCF_recs:
        key = get_variant_str(VCF_rec) # chr:pos:ref:alt
        variant_dict[key] = True
    return(variant_dict)

# From https://stackoverflow.com/questions/464864/how-to-get-all-possible-combinations-of-a-list-s-elements
# [A, B, C , D] -> [('A',), ('B',), ('C',), ('D',), ('A', 'B'), ('A', 'C'), ('A', 'D'), ('B', 'C'), ('B', 'D'), ('C', 'D'), ('A', 'B', 'C'), ('A', 'B', 'D'), ('A', 'C', 'D'), ('B', 'C', 'D'), ('A', 'B', 'C', 'D')]
def powerset(myset):
    subsets = []
    for L in range(0, len(myset)+1):
        for subset in itertools.combinations(myset, L):
            if len(subset) != 0:
                subsets.append(list(subset))
    return(subsets)

# From https://genome.ucsc.edu/goldenpath/help/hg19.chrom.sizes
# Also verified by checking the length of each sequence in human_g1k_v37_decoy.fasta - they agree
def get_hg19_chromosome_sizes():
    return({"1": 249250621, "2": 243199373, "3": 198022430, "4": 191154276, "5": 180915260, "6": 171115067, "7": 159138663, "8": 146364022, "9": 141213431, "10": 135534747, "11": 135006516, "12": 133851895, "13": 115169878, "14": 107349540, "15": 102531392, "16": 90354753, "17": 81195210, "18": 78077248, "19": 59128983, "20": 63025520, "21": 48129895, "22": 51304566, "X": 155270560, "Y": 59373566})

# From https://genome.ucsc.edu/goldenpath/help/hg38.chrom.sizes
# Also verified by checking the length of each sequence in Homo_sapiens_assembly38.fasta - they agree
def get_hg38_chromosome_sizes():
    return({"1": 248956422, "2": 242193529, "3": 198295559, "4": 190214555, "5": 181538259, "6": 170805979, "7": 159345973, "8": 145138636, "9": 138394717, "10": 133797422, "11": 135086622, "12": 133275309, "13": 114364328, "14": 107043718, "15": 101991189, "16": 90338345, "17": 83257441, "18": 80373285, "19": 58617616, "20": 64444167, "21": 46709983, "22": 50818468, "X": 156040895, "Y": 57227415})

def get_hg38_total_size():
    sum = 0
    hg38_chromosome_sizes = get_hg38_chromosome_sizes()
    for chrom in hg38_chromosome_sizes:
        sum += hg38_chromosome_sizes[chrom]
    return sum

def _add_string_fields_to_CNV_rec(this_CNV):

        this_CNV["filename_basename"] = os.path.basename(this_CNV["filename"])

        this_CNV["string_rep"] = this_CNV["chr"] + ":" + str(this_CNV["start"]) + "-" + str(this_CNV["end"])
        this_CNV["string_rep_extended"] = this_CNV["type"] + ":" + this_CNV["chr"] + ":" + str(this_CNV["start"]) + "-" + str(this_CNV["end"])
        this_CNV["string_rep_extended_size_caller"] = this_CNV["caller"] + ":" + this_CNV["type"] + ":" + this_CNV["chr"] + ":" + str(this_CNV["start"]) + "-" + str(this_CNV["end"]) + ":" + str(this_CNV["size"])
        this_CNV["string_rep_caller"] = this_CNV["caller"] + ":" + this_CNV["string_rep"]
        this_CNV["string_rep_caller_filename"] = this_CNV["filename_basename"] + ":" + this_CNV["caller"] + ":" + this_CNV["string_rep"]
        this_CNV["benchmark_list_rec"] = "{}:{}:{}:{}:{}:{}".format(this_CNV["chr"], this_CNV["start"], this_CNV["end"], this_CNV["size"], this_CNV["type"], this_CNV["caller"])
        return(this_CNV)

def is_header(potential_header):
    return(re.search("\sstart", potential_header, re.IGNORECASE))

# Combined [1,5] and [3,8] into [1,8] (etc.)
# Do this separately for each chromosome
# Input is a dictionary with start coordinate = "start", end coordinate = "end", and chromosome = "chr"; other elements of dictionary can be anything
def merge_overlapping_regions(regions):

    chromosome_regions = {} # Make separate lists, each containing a list from a given chromosome

    for region in regions:  # For each region, add this region to the list for the appropriate chromosome
        if region["chr"] not in chromosome_regions:
            chromosome_regions[region["chr"]] = []

        chromosome_regions[region["chr"]].append(region)

    merged_regions = [] # Maintain list of merged regions

    for chrom in sorted(chromosome_regions): # For each chromosome, sort the regions from the chromosome by start coordinate, then apply the merging algorithm
        regions = sorted(chromosome_regions[chrom], key=lambda region: region["start"])

        region_stack = []
        region_stack.append(regions[0])

        for i in range(1, len(regions)):
            top = region_stack[-1]

            if top["end"] < regions[i]["start"]:
                region_stack.append(regions[i])
            elif top["end"] < regions[i]["end"]:
                top["end"] = regions[i]["end"]
                region_stack[-1] = top

        merged_regions = merged_regions + region_stack

    return(merged_regions)


def depth_search(depth_filename, region):
    doc_file = open(depth_filename)
    index_file = open(depth_filename + ".idx")

    chrom, lower_limit, upper_limit = parse_region(region)

    mem_read_depth = 0
    mem_offset_pos = 0

    #Find offset number for seek()
    index_file.readline()
    for mline in index_file:
        mline = mline.split("\t")
        if chrom == mline[0] and int(mline[1]) <= lower_limit: # Get the offset that is one before the lower limit
            mem_offset_pos = int(mline[2])
        elif chrom == mline[0] and lower_limit < int(mline[1]):
            break

    #In case lower_limit == 1, or lower_limit < first offset bp pos
    if mem_offset_pos == 0:
        index_file.seek(0)
        for mline in index_file:
            mline = mline.split("\t")
            if chrom == mline[0]:
                mem_offset_pos = int(mline[2])
                break

    depth_list = []

    # Make list of read depths in requested region
    doc_file.seek(mem_offset_pos)
    for line in doc_file:
        docline_chr, docline_position, docline_depth = line.strip().split("\t")
        docline_chr = docline_chr.replace("chr", "")
        docline_position = int(docline_position)
        docline_depth = int(docline_depth)
        if chrom == docline_chr and (lower_limit <= docline_position <= upper_limit):
            depth_list.append(docline_depth)
        elif chrom == docline_chr and upper_limit < docline_position:
            break
        elif chrom != docline_chr:
            break

    return(depth_list)



def get_depth_lists(chrom, CNV_start, CNV_end, depth_filename, compute_left_left_right_right=False):
    region_template = "{0}:{1}-{2}"

    chrom_size = dict_from_file(depth_filename + ".chrinfo", header=True, value_type="int")

    if CNV_end < CNV_start:
        print("Invalid region entered (start coordinate {} is less than end coordinate {}). Please try again.".format(CNV_start, CNV_end))
        quit()
    if CNV_start < 1:
        print("Invalid region entered (start coordinate is less than 1). Please try again.")
        quit()
    if chrom not in chrom_size:
        print("Invalid chromomosome {} entered. Please try again.".format(chrom))
        quit()
    if CNV_end > chrom_size[chrom]:
        CNV_end = chrom_size[chrom]
        #print("Invalid region entered (end coordinate {} is greater than chromomosome {} size of {:d}). Please try again.".format(CNV_end, chrom, chrom_size[chrom]))
    if CNV_start > chrom_size[chrom]:
        return (None, None, None)

    CNV_size = CNV_end - CNV_start + 1

    left_flank_start = max(1, CNV_start - CNV_size)
    left_flank_end = CNV_start - 1
    left_flank_size = left_flank_end - left_flank_start + 1

    right_flank_start = CNV_end + 1
    right_flank_end = min(right_flank_start + CNV_size - 1, chrom_size[chrom])
    right_flank_size = right_flank_end - right_flank_start + 1

    left_flank_region = region_template.format(chrom, left_flank_start, left_flank_end)
    CNV_region = region_template.format(chrom, CNV_start, CNV_end)
    right_flank_region = region_template.format(chrom, right_flank_start, right_flank_end)

    left_flank_depths = depth_search(depth_filename, left_flank_region)
    CNV_depths = depth_search(depth_filename, CNV_region)
    right_flank_depths = depth_search(depth_filename, right_flank_region)

    depth_lists = {}
    depth_lists["left_flank_depths"] = left_flank_depths + [0] * (left_flank_size - len(left_flank_depths))
    depth_lists["CNV_depths"] = CNV_depths + [0] * (CNV_size - len(CNV_depths))
    depth_lists["right_flank_depths"] = right_flank_depths + [0] * (right_flank_size - len(right_flank_depths))

    if compute_left_left_right_right:  # For calculating read depth for short reads, if CNV is 251-300, compare with 151-200 and 351-400
        left_left_flank_start = max(1, CNV_start - 2*CNV_size)
        left_left_flank_end = max(1, left_flank_start - 1)
        left_left_flank_size = left_left_flank_end - left_left_flank_start + 1

        right_right_flank_start = min(right_flank_end + 1, chrom_size[chrom])
        right_right_flank_end = min(right_right_flank_start + CNV_size - 1, chrom_size[chrom])
        right_right_flank_size = right_right_flank_end - right_right_flank_start + 1

        left_left_flank_region = region_template.format(chrom, left_left_flank_start, left_left_flank_end)
        right_right_flank_region = region_template.format(chrom, right_right_flank_start, right_right_flank_end)

        left_left_flank_depths = depth_search(depth_filename, left_left_flank_region)
        right_right_flank_depths = depth_search(depth_filename, right_right_flank_region)

        depth_lists["left_left_flank_depths"] = left_left_flank_depths + [0] * (left_left_flank_size - len(left_left_flank_depths))
        depth_lists["right_right_flank_depths"] = right_right_flank_depths + [0] * (right_right_flank_size - len(right_right_flank_depths))

    return(depth_lists)

def parse_region(region): # Given a string in the form chr:start-end, return a three-element list containing chr, start, end
    chrom = region.split(":")[0].replace("chr", "").upper()
    start = int(region.split(":")[1].split("-")[0])
    end = int(region.split(":")[1].split("-")[1])

    return(chrom, start, end)

# For sorting purposes
def get_chromosome_num(chrom):
    if chrom == "X":
        return(23)
    elif chrom == "Y":
        return(24)
    elif chrom == "M":
        return(25)
    else:
        return(int(chrom))


def mean(list):
    try:
        mean = statistics.mean(list)
    except statistics.StatisticsError:
        mean = 0
    return(mean)

def stdev(list):
    try:
        stdev = statistics.stdev(list)
    except statistics.StatisticsError:
        stdev = 0
    return(stdev)


def generate_id(size=6, chars=string.ascii_lowercase + string.digits + string.ascii_uppercase):
    return(''.join(random.choice(chars) for _ in range(size)))

def is_int(str):
    try:
        str = int(str)
    except ValueError:
        return False
    return True

# Returns ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19', '20', '21', '22', 'X', 'Y']
def get_chromosomes():
    return([str(c) for c in list(range(1, 23)) + ["X", "Y"]])

# Returns ['chr1', 'chr2', 'chr3', 'chr4', 'chr5', 'chr6', 'chr7', 'chr8', 'chr9', 'chr10', 'chr11', 'chr12', 'chr13', 'chr14', 'chr15', 'chr16', 'chr17', 'chr18', 'chr19', 'chr20', 'chr21', 'chr22', 'chrX', 'chrY']
def get_chromosomes_chr_prefixed():
    return(["chr" + str(c) for c in list(range(1, 23)) + ["X", "Y"]])

# e.g., get_MSSNG_relation("1-0011-004", "Father")
def get_MSSNG_relation(ID, relation, filename="/Users/btrost/data/MSSNG/main_data_txt/MSSNG_sample_list_AGRE_annotated.txt"):
    file = open(filename)
    file.readline()
    lines = file.readlines()

    # Figure out the family ID for this individual
    for line in lines:
        fields = line.split("\t")
        if ID == fields[0]:
            family_ID = fields[1]
            break

    # Get the individual with the requested relation with the same family ID
    for line in lines:
        fields = line.split("\t")
        if fields[3] == relation and fields[1] == family_ID:
            return fields[0]

    return None

def reverse_complement(nt_sequence):
    return(nt_sequence[::-1].translate("".maketrans("ACGTacgt", "TGCAtgca")))

# Count the number of elements of group that are either a benchmark method or a CNV caller
def count_in_group(group, btype):
    in_group = {}
    for g in group:
        if g["btype"] == btype:
            in_group[g["caller"]] = 1
    return(len(in_group))

def count_microarray_in_group(group):
    in_group = {}
    for g in group:
        if g["caller"] in ["Affymetrix6.0", "Agilent24M", "Illumina1M", "NimbleGen42M"]:
            in_group[g["caller"]] = 1
    return(len(in_group))

# Sort groups first by number of predictions, and then by whether or not they are confirmed or not
# (Have at least one benchmark method in it)
def sort_groups(group):
    return((count_in_group(group, "prediction"), count_in_group(group, "benchmark")))

def get_size_bin(size, size_bins):
    for size_bin in size_bins:
        reg = re.search(r'(\d+),([\d.]+)', size_bin)
        lower = reg.group(1)
        upper = reg.group(2)

        if upper == "...":
            if size >= int(lower):
                return(size_bin)
        elif size >= int(lower) and size < int(upper):
            return(size_bin)

    return(None)


def column_as_array_from_file(filename, col_num=0, delimiter="\t", header=False):
    f = file_or_stdin(filename)
    array = []

    if header:
        f.readline()

    for line in f:
        array.append(line.rstrip("\n").split(delimiter)[col_num])
    return(array)

# Returns a filehandle either from the indicated filename, or if the indicated filename is "-", from standard input
def file_or_stdin(filename):
    if filename == "-":
        return(sys.stdin)
    else:
        return(open(filename))

def array_from_file(filename):
    f = file_or_stdin(filename)
    array = []

    for line in f:
        array.append(line.rstrip("\n"))
    return(array)

def same_repeat_unit(r1, r2):
    if len(r1) != len(r2):
        return(False)

    r2_extended = r2 + r2
    r2_extended_reverse_complement = reverse_complement(r2_extended)

    return(r1 in r2_extended or r1 in r2_extended_reverse_complement)

def same_orientation(r1, r2):
    if len(r1) != len(r2):
        return(False)

    r2_extended = r2 + r2
    return(r1 in r2_extended)


def get_set_from_file(filename, col=None, sep="\t", header=False):
    f = file_or_stdin(filename)
    d = set()
    if header:
        f.readline()
    for line in f:
        if col is not None:
            item = line.rstrip("\n").split(sep)[col]
        else:
            item = line.rstrip("\n")

        d.add(item)

    return(d)

def get_list_from_file(filename, col=None, sep="\t", header=False):
    f = file_or_stdin(filename)
    d = []
    if header:
        f.readline()
    for line in f:
        if col is not None:
            item = line.rstrip("\n").split(sep)[col]
        else:
            item = line.rstrip("\n")

        d.append(item)

    return(d)

# Read in a two column file, and make a dictionary where the
# key is the first column and the value is the second column
def dict_from_file(filename, delimiter="\t", key_col=0, value_col=1, header=False, value_type="string"):
    f = file_or_stdin(filename)

    if header:
        f.readline()

    d = {}

    for line in f:
        fields = line.split(delimiter)
        if value_type == "int":
            d[fields[key_col].rstrip("\n")] = int(fields[value_col].rstrip("\n"))
        elif value_type == "float":
            d[fields[key_col].rstrip("\n")] = float(fields[value_col].rstrip("\n"))
        else:
            d[fields[key_col].rstrip("\n")] = fields[value_col].rstrip("\n")

    return(d)

# Read in a two column file, and make a dictionary where the
# key is the first column and the value is the second column
def dict_of_lists_from_file(filename, delimiter="\t", key_col=0, value_col=1, header=False):
    f = file_or_stdin(filename)

    if header:
        f.readline()

    d = {}

    d = defaultdict(list)

    for line in f:
        fields = line.split(delimiter)
        d[fields[key_col].rstrip("\n")].append(fields[value_col].rstrip("\n"))

    return(d)

# Take a set of CNVs and make a dictionary whose keys are chromosomes and whose values are lists of CNVs in that chromosome
def CNVs_by_chromosome(CNVs):
    CNVs_by_chromosome = {}

    for c in get_chromosomes():
        CNVs_by_chromosome[c] = []

    for CNV in CNVs:
        if not CNV["chr"] in CNVs_by_chromosome:
            CNVs_by_chromosome[CNV["chr"]] = []

        CNVs_by_chromosome[CNV["chr"]].append(CNV)

    return(CNVs_by_chromosome)

def read_VCF(filename):
    file = file_or_stdin(filename)

    CNVs = []

    for line in file:
        this_CNV = {}
        line = line.rstrip("\n")
        this_CNV["full_line"] = line
        if re.match("#", line): # Ignore beginning comment lines
            continue
        fields = line.rstrip("\n").split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()
        this_CNV["start"] = int(fields[1])
        this_CNV["type"] = re.search(r"SVTYPE=(.+?);", line).group(1)

        if this_CNV["type"] == "BND":
            continue
        else:
            this_CNV["end"] = int(re.search(r"END=(\d+);", line).group(1))
            this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1

        this_CNV["caller"] = "NA"
        this_CNV["btype"] = "prediction"

        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        CNVs.append(this_CNV)

    return(CNVs, None)

def read_annotated(filename):
    file = file_or_stdin(filename)
    header = file.readline().rstrip("\n") # Discard header line
    CNVs = []

    for line in file:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["sample_ID"] = fields[0]
        this_CNV["chr"] = fields[1].replace("chr", "").upper()
        this_CNV["start"] = int(fields[2])
        this_CNV["end"] = int(fields[3])
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["type"] = fields[4]
        this_CNV["caller"] = "NA"

        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        CNVs.append(this_CNV)

    return(CNVs, header)


def read_CREST(filename):
    file = file_or_stdin(filename)
    CNVs = []

    for line in file:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()

        this_CNV["type"] = fields[8]

        if this_CNV["type"] == "INS":
            this_CNV["start"] = int(fields[5])
            this_CNV["end"] = int(fields[1])
        else:
            this_CNV["start"] = int(fields[1])
            this_CNV["end"] = int(fields[5])

        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["caller"] = "CREST"

        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        CNVs.append(this_CNV)

    return(CNVs, None)

# Read in an CNVnator file; returns a list of CNVs
# with various information (chromosome, start position, end position, size, type (deletion or duplication), and q0 value)
def read_CNVnator(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    if re.match("#", lines[0]):
        lines = lines[10:]

    CNVs = []

    for line in lines:

        if re.match("#", line): # Ignore beginning comment lines
            continue
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")
        if fields[0] == "duplication":
            this_CNV["type"] = "DUP"
        else:
            this_CNV["type"] = "DEL"

        this_CNV["chr"] = fields[1].split(":")[0].replace("chr", "").upper()
        this_CNV["start"] = int(fields[1].split(":")[1].split("-")[0])
        this_CNV["end"] = int(fields[1].split(":")[1].split("-")[1])
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["q0"] = float(fields[8])

        this_CNV["caller"] = "CNVnator"

        if this_CNV["chr"] == "M":
            continue

        this_CNV["filename"] = filename
        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        this_CNV["btype"] = "prediction"

        CNVs.append(this_CNV)

    return(CNVs, None)

# Read in an ERDS file; returns a list of CNVs
# with various information (chromosome, start position, end position, size, and type (deletion or duplication))
def read_ERDS(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    CNVs = []

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()
        this_CNV["start"] = int(fields[1])
        this_CNV["end"] = int(fields[2])
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["type"] = fields[4]
        this_CNV["caller"] = "ERDS"
        this_CNV["filename"] = filename

        if this_CNV["chr"] == "M":
            continue
        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        this_CNV["btype"] = "prediction"

        CNVs.append(this_CNV)

    return(CNVs, None)


def read_merged(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    CNVs = []
    header = lines[0].rstrip("\n")
    lines = lines[1:]

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()
        this_CNV["start"] = int(fields[1])
        this_CNV["end"] = int(fields[2])
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["type"] = fields[4]
        this_CNV["caller"] = "NA"

        if this_CNV["chr"] == "M":
            continue

        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)
        this_CNV["btype"] = "prediction"

        CNVs.append(this_CNV)

    return(CNVs, header)



# read in a BED file, which has the format
#chromosome start_pos   end_pos
def read_BED(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    if is_header(lines[0]):
        header = lines[0].rstrip("\n")
        lines = lines[1:]
    else:
        header = None

    CNVs = []

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        this_CNV["full_line_delimited"] = line.replace("\t", "***")
        fields = line.split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()
        this_CNV["start"] = int(fields[1])
        this_CNV["end"] = int(fields[2])
        this_CNV["size"] = int(float(fields[3]))
        if (len(fields) >= 5):
            this_CNV["type"] = fields[4]
        else:
            this_CNV["type"] = "N/A"

        if this_CNV["chr"] == "M":
            continue

        this_CNV["caller"] = filename
        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        this_CNV["btype"] = "prediction"
        CNVs.append(this_CNV)

    return(CNVs, header)

# read in one of Bhooma's ERDS+ files
def read_ERDS_plus(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    if is_header(lines[0]):
        header = lines[0].rstrip("\n")
        lines = lines[1:]
    else:
        header = None

    CNVs = []

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        this_CNV["full_line_delimited"] = line.replace("\t", "***")
        fields = line.split("\t")

        this_CNV["chr"] = fields[1].replace("chr", "").upper()
        this_CNV["start"] = int(fields[2])
        this_CNV["end"] = int(fields[3])
        this_CNV["type"] = fields[4]
        this_CNV["size"] = int(fields[5])

        this_CNV["caller"] = filename
        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        this_CNV["btype"] = "prediction"
        CNVs.append(this_CNV)

    return(CNVs, header)


# read in a "filename BED" file (my own made-up name), which has the format
#filename chromosome start_pos   end_pos ...
def read_filename_BED(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    if is_header(lines[0]):
        header = lines[0].rstrip("\n")
        lines = lines[1:]
    else:
        header = None

    CNVs = []

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["BAM_filename"] = fields[0]
        this_CNV["chr"] = fields[1].replace("chr", "").upper()
        this_CNV["start"] = int(fields[2])
        this_CNV["end"] = int(fields[3])
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        if (len(fields) >= 6):
            this_CNV["type"] = fields[5]
        else:
            this_CNV["type"] = "N/A"

        if this_CNV["chr"] == "M":
            continue

        this_CNV["caller"] = "NA"
        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        this_CNV["btype"] = "prediction"
        CNVs.append(this_CNV)

    return(CNVs, header)

def read_filtered_CNV_benchmark(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    CNVs = []

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()
        this_CNV["start"] = int(fields[1])
        this_CNV["end"] = int(fields[2])
        this_CNV["size"] = int(fields[3])
        if (len(fields) >= 5):
            this_CNV["type"] = fields[4]
        else:
            this_CNV["type"] = "N/A"

        if this_CNV["chr"] == "M":
            continue

        this_CNV["caller"] = "NA"
        this_CNV["filename"] = filename

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        this_CNV["btype"] = "prediction"

        CNVs.append(this_CNV)

    return(CNVs, None)

# read in a benchmark file, which has the format
#chromosome start_pos   end_pos length  type    technology
def read_benchmark(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()

    CNVs = []
    if is_header(lines[0]):
        header = lines[0].rstrip("\n")
        lines = lines[1:]
    else:
        header = None

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()
        this_CNV["start"] = int(fields[1])
        this_CNV["end"] = int(fields[2])
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["type"] = fields[4]
        this_CNV["caller"] = fields[5]
        this_CNV["filename"] = filename

        if this_CNV["chr"] == "M":
            continue

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        this_CNV["btype"] = "benchmark"

        CNVs.append(this_CNV)

    return(CNVs, header)

def read_Miriam_CNV_file(filename):
    CNV_file = open(filename)
    CNV_file.readline()

    # Read in the CNVs
    CNVs = {}
    for line in CNV_file:
        this_CNV = {}
        this_CNV["full_line"] = line[-1]
        this_CNV["sample"], this_CNV["chr"], this_CNV["start"], this_CNV["end"], this_CNV["type"] = line.rstrip("\n").split("\t")
        this_CNV["chr"] = this_CNV["chr"].upper()
        this_CNV["start"] = int(this_CNV["start"])
        this_CNV["end"] = int(this_CNV["end"])


        if this_CNV["chr"] == "M":
            continue

        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1

        if this_CNV["sample"] not in CNVs:
            CNVs[this_CNV["sample"]] = []

        this_CNV["filename"] = filename

        this_CNV["caller"] = "NA"

        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        CNVs[this_CNV["sample"]].append(this_CNV)

    return CNVs

# read in one of Zhuozhi's combined CNV files, which has the format
#chromosome start_pos   end_pos length  DEL/DUP important_info  caller_name caller_all_info length_of_ovelap_with_benchmark benchmark_matches   length_of_overlap_with_repeats  list_of_repeats_overlapped
def read_Zhuozhi(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()
    CNVs = []

    if is_header(lines[0]):
        header = lines[0].rstrip("\n")
        lines = lines[1:]
    else:
        header = None

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["chr"] = fields[0].replace("chr", "").upper()
        this_CNV["start"] = int(float(fields[1]))
        this_CNV["end"] = int(float(fields[2]))
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["type"] = fields[4]
        this_CNV["important_caller_info"] = fields[5]
        this_CNV["caller"] = fields[6]
        this_CNV["all_caller_info"] = fields[7]
        #this_CNV["benchmark_overlap_lengths"] = [int(x) for x in fields[8].split("|")]
        #this_CNV["benchmark_overlap_details"] = fields[9].split("|")
        #this_CNV["repeat_overlap_lengths"] = [int(x) for x in fields[10].split("|")]
        #this_CNV["repeat_overlap_details"] = fields[11].split(":")

        if this_CNV["chr"] == "M":
            continue

        this_CNV["benchmark_overlap_details_parsed"] = {}
        this_CNV["filename"] = filename
        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)
        this_CNV["btype"] = "prediction"

        this_CNV["string_rep_caller_filename_callerinfo"] = "{}|{}".format(this_CNV["string_rep_caller_filename"], this_CNV["all_caller_info"])

        CNVs.append(this_CNV)

    return(CNVs, header)

# read a CNV file that has the following format:
# PGPC-0001	1	664001	727000	DUP
def read_different_samples(filename):
    file = file_or_stdin(filename)
    lines = file.readlines()
    CNVs = []

    for line in lines:
        line = line.rstrip("\n")
        this_CNV = {}
        this_CNV["full_line"] = line
        fields = line.split("\t")

        this_CNV["sample"] = fields[0]
        this_CNV["chr"] = fields[1].replace("chr", "").upper()
        this_CNV["start"] = int(float(fields[2]))
        this_CNV["end"] = int(float(fields[3]))
        this_CNV["size"] = this_CNV["end"] - this_CNV["start"] + 1
        this_CNV["type"] = fields[4]
        this_CNV["filename"] = filename
        this_CNV["caller"] = "NA"
        this_CNV = _add_string_fields_to_CNV_rec(this_CNV)

        CNVs.append(this_CNV)

    return(CNVs, None)

# Determines the format of a file that contains CNVs.
# Returns a two-item dictionary, with the "string" element being a string representation of the type determined, and "func" being the function
# that reads a file in that particular format
def determine_CNV_format(filename):
    file = open(filename)
    data = file.read()
    num_cols = len(data.split("\t"))

    result = {}

    if (re.search("NotApplicable", data) or re.search("ERDS\+CNVnator", data)) and re.search("SIMPLE", data):
        result["string"] = "Zhuozhi"
        result["function"] = read_Zhuozhi
    elif re.search("CNVN_Boundaries", data):
        result["string"] = "annotated"
        result["function"] = read_annotated
    elif re.search("\tCTX\t", data):
        result["string"] = "CREST"
        result["function"] = read_CREST
    elif re.search("source=LUMPY", data) or re.search("MantaBND", data) or re.search("EMBL.DELLY", data):
        result["string"] = "VCF"
        result["function"] = read_VCF
    elif re.search("Number of methods detecting this CNV", data):
        result["string"] = "Merged (output from CNV_overlap.py)"
        result["function"] = read_merged
    elif re.search("SplitRead", data) and re.search("hetero", data):
        result["string"] = "VenterSplitRead"
        result["function"] = read_VenterSplitRead
    elif re.search("Affymetrix6.0|Affy6|Agilent24M|Agilent.24M|Illumina1M|Illumina.1M|MatePair|NimbleGen42M|NimbleGen.42M|SplitRead|split.read|assembly.comparison|cg.cnv|cg.sp|cg.mp|Personalis_1000_Genomes_deduplicated", data):
        result["string"] = "benchmark"
        result["function"] = read_benchmark
    elif re.search("duplication", data):
        result["string"] = "CNVnator"
        result["function"] = read_CNVnator
    elif num_cols == 8:
        result["string"] = "ERDS"
        result["function"] = read_ERDS
    else:
        result["string"] = "BED"
        result["function"] = read_BED
    return(result)

def read_autodetect(filename, print_format=False):
    CNV_file_format = determine_CNV_format(filename)

    if print_format:
        print("Detected format is " + CNV_file_format["string"])
    return(CNV_file_format["function"](filename))

# Find the size distribution of a list of CNVs (or potentially other genomic regions)
def SV_size_dist(SVs, size_bins):
    size_dist = defaultdict(lambda: Counter())

    for SV in SVs:
        s = get_size_bin(SV["size"], size_bins)
        if s:
            size_dist[SV["type"]][s] += 1
            size_dist[SV["type"]]["Total"] += 1

    return(size_dist)

determine_SV_format = determine_CNV_format # Alias for SV paper

def num_lines_in_file(file):
    return sum(1 for line in file)

# From http://stackoverflow.com/questions/2953967/built-in-function-for-computing-overlap-in-python
def get_overlap(a, b, zero_based=False): # a = [5,10] b = [7,12]
    if zero_based:
        a[0] = a[0] + 1
        b[0] = b[0] + 1
    return(max(0, min(a[1], b[1]) - max(a[0], b[0]) + 1))

# https://stackoverflow.com/questions/2130016/splitting-a-list-into-n-parts-of-approximately-equal-length/37414115
# NOT the accepted answer!
def split_array_into_equal_parts(l, num_parts):
    k, m = divmod(len(l), num_parts)
    return list((l[i * k + min(i, m):(i + 1) * k + min(i + 1, m)] for i in range(num_parts)))

def get_overlap_region(CNV1, CNV2):
    return get_overlap([CNV1["start"], CNV1["end"]], [CNV2["start"], CNV2["end"]])

def fifty_percent_reciprocal_overlap(CNV1, CNV2, ignore_direction=False):
    if CNV1["chr"] != CNV2["chr"]:
        return(False)

    if CNV1["type"] != CNV2["type"] and not ignore_direction:
        return(False)

    if CNV1["type"] == "INS" and CNV2["type"] == "INS":
        return(abs(CNV1["start"] - CNV2["start"]) <= 50 and CNV1["size"] >= CNV2["size"]/2 and CNV2["size"] >= CNV1["size"]/2)

    if CNV1["type"] == "INS" or CNV1["type"] == "BND":
        return(abs(CNV1["start"] - CNV2["start"]) <= 50)

    overlap = get_overlap([CNV1["start"], CNV1["end"]], [CNV2["start"], CNV2["end"]])
    return(overlap >= (CNV1["size"] / 2) and overlap >= (CNV2["size"] / 2))

def avg_reciprocal_overlap(CNV1, CNV2):
    overlap = get_overlap([CNV1["start"], CNV1["end"]], [CNV2["start"], CNV2["end"]])

    return(((overlap / CNV1["size"] * 100) + (overlap / CNV2["size"] * 100))/2)

def reciprocal_overlap(CNV1, CNV2, percent):
    if CNV1["chr"] != CNV2["chr"]:
        return(False)

    overlap = get_overlap([CNV1["start"], CNV1["end"]], [CNV2["start"], CNV2["end"]])
    return(overlap >= (CNV1["size"] * percent/100) and overlap >= (CNV2["size"] * percent/100))

def any_overlap(CNV1, CNV2, ignore_direction=False):
    if CNV1["chr"] != CNV2["chr"]:
        return(False)

    if CNV1["type"] != CNV2["type"] and not ignore_direction:
        return(False)

    overlap = get_overlap([CNV1["start"], CNV1["end"]], [CNV2["start"], CNV2["end"]])
    return(overlap > 0)

def debugprint(s, debug_mode):
    if debug_mode:
        print(s)

def median_size_of_CNVs_in_group(group):
    size_list = []
    for g in group:
        size_list.append(g["size"])

    return(statistics.median(size_list))

def get_RLCR_trees(RLCR_definition_filename):
    chromosomes = get_chromosomes()
    trees = {}
    RLCR_definition_file = open(RLCR_definition_filename)

    # Make a new interval tree for each chromosome
    for chrom in chromosomes:
        trees[chrom] = IntervalTree()

    # Read RLCR definition file into interval tree data structure
    for line in RLCR_definition_file:
        chrom, start, end, repeat_type = line.rstrip("\n").split("\t")

        if chrom == "Chromosome":
            continue

        chrom = chrom.replace("chr", "")
        start = int(start)
        end = int(end)
        trees[chrom][start:end+1] = repeat_type # Insert this interval into the tree
    RLCR_definition_file.close()

    # Make a deep copy of the tree so that we can merge overlapping intervals in the copied one
    trees_merged = copy.deepcopy(trees)
    for chrom in chromosomes:
        trees_merged[chrom].merge_overlaps()

    return(trees, trees_merged)


median_size_of_SVs_in_group = median_size_of_CNVs_in_group

read_functions = {}
read_functions["VCF"] = read_VCF
read_functions["annotated"] = read_annotated
read_functions["CREST"] = read_CREST
read_functions["CNVnator"] = read_CNVnator
read_functions["ERDS"] = read_ERDS
read_functions["merged"] = read_merged
read_functions["BED"] = read_BED
read_functions["filename_BED"] = read_filename_BED
read_functions["filtered_CNV_benchmark"] = read_filtered_CNV_benchmark
read_functions["benchmark"] = read_benchmark
read_functions["Miriam_CNV_file"] = read_Miriam_CNV_file
read_functions["Zhuozhi"] = read_Zhuozhi
read_functions["ERDS_plus"] = read_ERDS_plus
read_functions["autodetect"] = read_autodetect
